"""
utils/transcript_reader.py

Business logic for extracting raw text from uploaded meeting transcript
files. This module has NO Streamlit code — it only takes file-like
objects and returns plain text or raises a controlled TranscriptReadError.

Supported formats:
    - .txt
    - .docx
    - .pdf
"""

from dataclasses import dataclass
from io import BytesIO

import docx
from pypdf import PdfReader


class TranscriptReadError(Exception):
    """Raised when a transcript file cannot be read or parsed."""
    pass


@dataclass
class ReadResult:
    """Result of attempting to read a transcript file."""
    text: str
    filename: str
    file_type: str
    char_count: int


SUPPORTED_EXTENSIONS = {"txt", "docx", "pdf"}


def get_file_extension(filename: str) -> str:
    """Returns the lowercase extension of a filename, without the dot."""
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def read_txt(file_bytes: bytes) -> str:
    """
    Decodes raw bytes from a .txt file into text.
    Tries UTF-8 first, falls back to latin-1 for legacy files.
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as exc:
            raise TranscriptReadError(
                "Could not decode the .txt file. It may use an "
                "unsupported text encoding."
            ) from exc


def read_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a .docx file, including paragraphs and table cells.
    """
    try:
        document = docx.Document(BytesIO(file_bytes))
    except Exception as exc:
        raise TranscriptReadError(
            "Could not open the .docx file. It may be corrupted or "
            "not a valid Word document."
        ) from exc

    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    text = "\n".join(parts)

    if not text.strip():
        raise TranscriptReadError(
            "The .docx file appears to contain no readable text."
        )

    return text


def read_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a .pdf file, page by page.
    """
    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception as exc:
        raise TranscriptReadError(
            "Could not open the .pdf file. It may be corrupted, "
            "password-protected, or not a valid PDF."
        ) from exc

    if reader.is_encrypted:
        raise TranscriptReadError(
            "This PDF is password-protected. Please upload an "
            "unprotected version."
        )

    parts = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            parts.append(page_text)

    text = "\n".join(parts)

    if not text.strip():
        raise TranscriptReadError(
            "No extractable text was found in this PDF. It may be a "
            "scanned/image-based PDF, which requires OCR (not yet "
            "supported)."
        )

    return text


def read_transcript_file(filename: str, file_bytes: bytes) -> ReadResult:
    """
    Main entry point: detects file type from filename extension and
    dispatches to the correct reader. Raises TranscriptReadError with
    a user-friendly message on any failure.
    """
    if not file_bytes:
        raise TranscriptReadError("The uploaded file is empty.")

    extension = get_file_extension(filename)

    if extension not in SUPPORTED_EXTENSIONS:
        raise TranscriptReadError(
            f"Unsupported file type '.{extension}'. "
            f"Please upload a .txt, .docx, or .pdf file."
        )

    if extension == "txt":
        text = read_txt(file_bytes)
    elif extension == "docx":
        text = read_docx(file_bytes)
    elif extension == "pdf":
        text = read_pdf(file_bytes)
    else:
        # Defensive fallback; should never reach here given the check above
        raise TranscriptReadError(f"Unsupported file type '.{extension}'.")

    if not text.strip():
        raise TranscriptReadError(
            "No readable text could be extracted from this file."
        )

    return ReadResult(
        text=text,
        filename=filename,
        file_type=extension,
        char_count=len(text),
    )