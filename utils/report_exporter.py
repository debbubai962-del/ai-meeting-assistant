"""
utils/report_exporter.py

Business logic for exporting a structured meeting analysis result into
TXT, Markdown, DOCX, and PDF formats. No Streamlit code here.
"""

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
)


def _section_lines(title: str, items: list) -> list:
    """Returns a list of plain-text lines for a bulleted section."""
    lines = [title.upper(), "-" * len(title)]
    if not items:
        lines.append("(None identified)")
    else:
        for item in items:
            lines.append(f"  - {item}")
    lines.append("")
    return lines


def build_report_filename(metadata: dict, extension: str) -> str:
    """Builds a clean, safe filename from the meeting title and date."""
    title = (metadata.get("meeting_title") or "meeting_report").strip()
    safe_title = "".join(c if c.isalnum() or c in (" ", "-", "_") else "" for c in title)
    safe_title = safe_title.replace(" ", "_")[:60] or "meeting_report"
    date_str = metadata.get("meeting_date", datetime.now().strftime("%Y-%m-%d"))
    return f"{safe_title}_{date_str}.{extension}"


# -----------------------------------------------------------------------
# Plain text
# -----------------------------------------------------------------------

def build_txt(data: dict, metadata: dict) -> str:
    """Builds a plain-text version of the report."""
    lines = []
    lines.append("AI BUSINESS MEETING ASSISTANT — MEETING REPORT")
    lines.append("=" * 50)
    lines.append("")
    lines.append(f"Title: {metadata.get('meeting_title', 'Not specified')}")
    lines.append(f"Type: {metadata.get('meeting_type', 'Not specified')}")
    lines.append(f"Date: {metadata.get('meeting_date', 'Not specified')}")
    lines.append(f"Duration: {metadata.get('duration_minutes', 'Not specified')} minutes")
    lines.append(f"Organizer: {metadata.get('organizer') or 'Not specified'}")
    lines.append(f"Department: {metadata.get('department') or 'Not specified'}")
    participants = ", ".join(metadata.get("participants", [])) or "Not specified"
    lines.append(f"Participants: {participants}")
    lines.append(f"Sentiment: {data.get('meeting_sentiment', 'Neutral')}")
    lines.append(f"Priority: {data.get('overall_priority_level', 'Medium')}")
    lines.append("")

    lines.append("EXECUTIVE SUMMARY")
    lines.append("-" * 17)
    lines.append(data.get("executive_summary") or "Not available.")
    lines.append("")

    lines.append("DETAILED SUMMARY")
    lines.append("-" * 16)
    lines.append(data.get("detailed_summary") or "Not available.")
    lines.append("")

    lines.extend(_section_lines("Decisions Taken", data.get("decisions_taken", [])))

    lines.append("ACTION ITEMS")
    lines.append("-" * 12)
    action_items = data.get("action_items", [])
    if not action_items:
        lines.append("(None identified)")
    else:
        for idx, item in enumerate(action_items, start=1):
            lines.append(
                f"  {idx}. {item.get('task', '')} "
                f"[Owner: {item.get('owner') or 'N/A'}, "
                f"Team: {item.get('team') or 'N/A'}, "
                f"Deadline: {item.get('deadline') or 'N/A'}, "
                f"Priority: {item.get('priority', 'Medium')}]"
            )
    lines.append("")

    lines.extend(_section_lines("Risks", data.get("risks", [])))
    lines.extend(_section_lines("Opportunities", data.get("opportunities", [])))
    lines.extend(_section_lines("Blockers", data.get("blockers", [])))
    lines.extend(_section_lines("Open Questions", data.get("open_questions", [])))
    lines.extend(_section_lines("KPI Mentions", data.get("kpi_mentions", [])))
    lines.extend(_section_lines("Budget Mentions", data.get("budget_mentions", [])))
    lines.extend(_section_lines("Timeline Mentions", data.get("timeline_mentions", [])))
    lines.extend(_section_lines("Compliance Issues", data.get("compliance_issues", [])))
    lines.extend(_section_lines("Escalations", data.get("escalations", [])))
    lines.extend(_section_lines("Key Business Insights", data.get("key_business_insights", [])))
    lines.extend(_section_lines("Suggested Next Steps", data.get("suggested_next_steps", [])))
    lines.extend(_section_lines("Next Meeting Agenda", data.get("next_meeting_agenda", [])))

    email = data.get("follow_up_email", {})
    lines.append("FOLLOW-UP EMAIL")
    lines.append("-" * 15)
    lines.append(f"Subject: {email.get('subject', '')}")
    lines.append("")
    lines.append(email.get("body", ""))
    lines.append("")

    return "\n".join(lines)


# -----------------------------------------------------------------------
# Markdown
# -----------------------------------------------------------------------

def build_markdown(data: dict, metadata: dict) -> str:
    """Builds a Markdown version of the report."""
    md = []
    md.append(f"# {metadata.get('meeting_title', 'Meeting Report')}")
    md.append("")
    md.append(f"**Type:** {metadata.get('meeting_type', 'Not specified')}  ")
    md.append(f"**Date:** {metadata.get('meeting_date', 'Not specified')}  ")
    md.append(f"**Duration:** {metadata.get('duration_minutes', 'Not specified')} minutes  ")
    md.append(f"**Organizer:** {metadata.get('organizer') or 'Not specified'}  ")
    md.append(f"**Department:** {metadata.get('department') or 'Not specified'}  ")
    participants = ", ".join(metadata.get("participants", [])) or "Not specified"
    md.append(f"**Participants:** {participants}  ")
    md.append(f"**Sentiment:** {data.get('meeting_sentiment', 'Neutral')}  ")
    md.append(f"**Priority:** {data.get('overall_priority_level', 'Medium')}")
    md.append("")

    md.append("## Executive Summary")
    md.append(data.get("executive_summary") or "_Not available._")
    md.append("")

    md.append("## Detailed Summary")
    md.append(data.get("detailed_summary") or "_Not available._")
    md.append("")

    def bullet_section(title: str, items: list) -> None:
        md.append(f"## {title}")
        if not items:
            md.append("_None identified._")
        else:
            for item in items:
                md.append(f"- {item}")
        md.append("")

    bullet_section("Decisions Taken", data.get("decisions_taken", []))

    md.append("## Action Items")
    action_items = data.get("action_items", [])
    if not action_items:
        md.append("_None identified._")
    else:
        md.append("| # | Task | Owner | Team | Deadline | Priority |")
        md.append("|---|------|-------|------|----------|----------|")
        for idx, item in enumerate(action_items, start=1):
            md.append(
                f"| {idx} | {item.get('task', '')} | {item.get('owner') or '—'} "
                f"| {item.get('team') or '—'} | {item.get('deadline') or '—'} "
                f"| {item.get('priority', 'Medium')} |"
            )
    md.append("")

    bullet_section("Risks", data.get("risks", []))
    bullet_section("Opportunities", data.get("opportunities", []))
    bullet_section("Blockers", data.get("blockers", []))
    bullet_section("Open Questions", data.get("open_questions", []))
    bullet_section("KPI Mentions", data.get("kpi_mentions", []))
    bullet_section("Budget Mentions", data.get("budget_mentions", []))
    bullet_section("Timeline Mentions", data.get("timeline_mentions", []))
    bullet_section("Compliance Issues", data.get("compliance_issues", []))
    bullet_section("Escalations", data.get("escalations", []))
    bullet_section("Key Business Insights", data.get("key_business_insights", []))
    bullet_section("Suggested Next Steps", data.get("suggested_next_steps", []))
    bullet_section("Next Meeting Agenda", data.get("next_meeting_agenda", []))

    email = data.get("follow_up_email", {})
    md.append("## Follow-Up Email")
    md.append(f"**Subject:** {email.get('subject', '')}")
    md.append("")
    md.append(email.get("body", "") or "_Not available._")
    md.append("")

    return "\n".join(md)


# -----------------------------------------------------------------------
# DOCX
# -----------------------------------------------------------------------

def build_docx(data: dict, metadata: dict) -> bytes:
    """Builds a Word (.docx) version of the report, returned as bytes."""
    document = Document()

    title = document.add_heading(metadata.get("meeting_title", "Meeting Report"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_para = document.add_paragraph()
    participants = ", ".join(metadata.get("participants", [])) or "Not specified"
    meta_lines = [
        f"Type: {metadata.get('meeting_type', 'Not specified')}",
        f"Date: {metadata.get('meeting_date', 'Not specified')}",
        f"Duration: {metadata.get('duration_minutes', 'Not specified')} minutes",
        f"Organizer: {metadata.get('organizer') or 'Not specified'}",
        f"Department: {metadata.get('department') or 'Not specified'}",
        f"Participants: {participants}",
        f"Sentiment: {data.get('meeting_sentiment', 'Neutral')}",
        f"Priority: {data.get('overall_priority_level', 'Medium')}",
    ]
    for line in meta_lines:
        run = meta_para.add_run(line + "\n")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(data.get("executive_summary") or "Not available.")

    document.add_heading("Detailed Summary", level=1)
    document.add_paragraph(data.get("detailed_summary") or "Not available.")

    def add_bullet_section(title: str, items: list) -> None:
        document.add_heading(title, level=1)
        if not items:
            document.add_paragraph("None identified.")
        else:
            for item in items:
                document.add_paragraph(str(item), style="List Bullet")

    add_bullet_section("Decisions Taken", data.get("decisions_taken", []))

    document.add_heading("Action Items", level=1)
    action_items = data.get("action_items", [])
    if not action_items:
        document.add_paragraph("None identified.")
    else:
        table = document.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for idx, heading in enumerate(["Task", "Owner", "Team", "Deadline", "Priority"]):
            header_cells[idx].text = heading
        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("task", "")
            row_cells[1].text = item.get("owner") or "—"
            row_cells[2].text = item.get("team") or "—"
            row_cells[3].text = item.get("deadline") or "—"
            row_cells[4].text = item.get("priority", "Medium")

    add_bullet_section("Risks", data.get("risks", []))
    add_bullet_section("Opportunities", data.get("opportunities", []))
    add_bullet_section("Blockers", data.get("blockers", []))
    add_bullet_section("Open Questions", data.get("open_questions", []))
    add_bullet_section("KPI Mentions", data.get("kpi_mentions", []))
    add_bullet_section("Budget Mentions", data.get("budget_mentions", []))
    add_bullet_section("Timeline Mentions", data.get("timeline_mentions", []))
    add_bullet_section("Compliance Issues", data.get("compliance_issues", []))
    add_bullet_section("Escalations", data.get("escalations", []))
    add_bullet_section("Key Business Insights", data.get("key_business_insights", []))
    add_bullet_section("Suggested Next Steps", data.get("suggested_next_steps", []))
    add_bullet_section("Next Meeting Agenda", data.get("next_meeting_agenda", []))

    email = data.get("follow_up_email", {})
    document.add_heading("Follow-Up Email", level=1)
    subject_para = document.add_paragraph()
    subject_run = subject_para.add_run(f"Subject: {email.get('subject', '')}")
    subject_run.bold = True
    document.add_paragraph(email.get("body", "") or "Not available.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


# -----------------------------------------------------------------------
# PDF
# -----------------------------------------------------------------------

def build_pdf(data: dict, metadata: dict) -> bytes:
    """Builds a PDF version of the report, returned as bytes."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], fontSize=20, spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "ReportHeading", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6,
        textColor=colors.HexColor("#1E3A8A"),
    )
    meta_style = ParagraphStyle(
        "MetaText", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#555555"),
    )
    body_style = styles["Normal"]

    story = []

    story.append(Paragraph(metadata.get("meeting_title", "Meeting Report"), title_style))

    participants = ", ".join(metadata.get("participants", [])) or "Not specified"
    meta_text = (
        f"Type: {metadata.get('meeting_type', 'Not specified')}<br/>"
        f"Date: {metadata.get('meeting_date', 'Not specified')}<br/>"
        f"Duration: {metadata.get('duration_minutes', 'Not specified')} minutes<br/>"
        f"Organizer: {metadata.get('organizer') or 'Not specified'}<br/>"
        f"Department: {metadata.get('department') or 'Not specified'}<br/>"
        f"Participants: {participants}<br/>"
        f"Sentiment: {data.get('meeting_sentiment', 'Neutral')} &nbsp;|&nbsp; "
        f"Priority: {data.get('overall_priority_level', 'Medium')}"
    )
    story.append(Paragraph(meta_text, meta_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(data.get("executive_summary") or "Not available.", body_style))

    story.append(Paragraph("Detailed Summary", heading_style))
    story.append(Paragraph(data.get("detailed_summary") or "Not available.", body_style))

    def add_bullet_section(title: str, items: list) -> None:
        story.append(Paragraph(title, heading_style))
        if not items:
            story.append(Paragraph("None identified.", body_style))
        else:
            bullet_items = [ListItem(Paragraph(str(i), body_style)) for i in items]
            story.append(ListFlowable(bullet_items, bulletType="bullet"))

    add_bullet_section("Decisions Taken", data.get("decisions_taken", []))

    story.append(Paragraph("Action Items", heading_style))
    action_items = data.get("action_items", [])
    if not action_items:
        story.append(Paragraph("None identified.", body_style))
    else:
        table_data = [["Task", "Owner", "Team", "Deadline", "Priority"]]
        for item in action_items:
            table_data.append([
                item.get("task", ""), item.get("owner") or "—",
                item.get("team") or "—", item.get("deadline") or "—",
                item.get("priority", "Medium"),
            ])
        action_table = Table(table_data, repeatRows=1, colWidths=[1.8 * inch, 1 * inch, 1 * inch, 1 * inch, 0.8 * inch])
        action_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(action_table)

    add_bullet_section("Risks", data.get("risks", []))
    add_bullet_section("Opportunities", data.get("opportunities", []))
    add_bullet_section("Blockers", data.get("blockers", []))
    add_bullet_section("Open Questions", data.get("open_questions", []))
    add_bullet_section("KPI Mentions", data.get("kpi_mentions", []))
    add_bullet_section("Budget Mentions", data.get("budget_mentions", []))
    add_bullet_section("Timeline Mentions", data.get("timeline_mentions", []))
    add_bullet_section("Compliance Issues", data.get("compliance_issues", []))
    add_bullet_section("Escalations", data.get("escalations", []))
    add_bullet_section("Key Business Insights", data.get("key_business_insights", []))
    add_bullet_section("Suggested Next Steps", data.get("suggested_next_steps", []))
    add_bullet_section("Next Meeting Agenda", data.get("next_meeting_agenda", []))

    email = data.get("follow_up_email", {})
    story.append(Paragraph("Follow-Up Email", heading_style))
    story.append(Paragraph(f"<b>Subject:</b> {email.get('subject', '')}", body_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph((email.get("body", "") or "Not available.").replace("\n", "<br/>"), body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()